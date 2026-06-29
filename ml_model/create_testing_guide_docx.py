from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
import os

script_dir = os.path.dirname(os.path.abspath(__file__))

NAVY = "17365D"
BLUE = "2E74B5"
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F2F4F7"
GOLD = "7A5A00"
LIGHT_GOLD = "FFF4CC"
RED = "9B1C1C"
LIGHT_RED = "FDECEC"
GREEN = "1E6B3A"
WHITE = "FFFFFF"
BLACK = "000000"
MUTED = "666666"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=100, start=120, bottom=100, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_table_borders(table, color="B7C9DC", size="6"):
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        element = borders.find(qn(f"w:{edge}"))
        if element is None:
            element = OxmlElement(f"w:{edge}")
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_font(run, name="Calibri", size=11, bold=False, italic=False, color=BLACK):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    run.font.color.rgb = RGBColor.from_string(color)


def add_page_number(paragraph):
    run = paragraph.add_run()
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = "PAGE"
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char1)
    run._r.append(instr_text)
    run._r.append(fld_char2)
    set_font(run, size=9, color=MUTED)


def add_heading(doc, text, level=1):
    p = doc.add_paragraph(style=f"Heading {level}")
    p.add_run(text)
    return p


def add_body(doc, text, bold_prefix=None):
    p = doc.add_paragraph(style="Normal")
    if bold_prefix and text.startswith(bold_prefix):
        r1 = p.add_run(bold_prefix)
        set_font(r1, bold=True)
        r2 = p.add_run(text[len(bold_prefix):])
        set_font(r2)
    else:
        r = p.add_run(text)
        set_font(r)
    return p


def add_bullet(doc, text, level=0):
    style = "List Bullet" if level == 0 else "List Bullet 2"
    p = doc.add_paragraph(style=style)
    r = p.add_run(text)
    set_font(r)
    return p


def add_number(doc, text):
    p = doc.add_paragraph(style="List Number")
    r = p.add_run(text)
    set_font(r)
    return p


def add_formula(doc, formula, explanation=None):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(5)
    p.paragraph_format.space_after = Pt(5)
    r = p.add_run(formula)
    set_font(r, name="Cambria Math", size=11.5, bold=True, color=NAVY)
    if explanation:
        p2 = doc.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p2.paragraph_format.space_after = Pt(6)
        r2 = p2.add_run(explanation)
        set_font(r2, size=9.5, italic=True, color=MUTED)


def add_callout(doc, title, body, fill=LIGHT_BLUE, title_color=NAVY):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.columns[0].width = Inches(6.3)
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    set_cell_margins(cell, top=150, start=180, bottom=150, end=180)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(title)
    set_font(r, bold=True, color=title_color)
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_after = Pt(0)
    r2 = p2.add_run(body)
    set_font(r2, size=10.5)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)


def add_table(doc, headers, rows, widths=None, font_size=9.5):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    set_table_borders(table)
    hdr = table.rows[0]
    set_repeat_table_header(hdr)
    for i, header in enumerate(headers):
        cell = hdr.cells[i]
        set_cell_shading(cell, NAVY)
        set_cell_margins(cell)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(str(header))
        set_font(r, size=font_size, bold=True, color=WHITE)
    for row_idx, row_data in enumerate(rows):
        cells = table.add_row().cells
        for i, value in enumerate(row_data):
            cell = cells[i]
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            if row_idx % 2 == 1:
                set_cell_shading(cell, LIGHT_GRAY)
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT if i == 0 else WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_after = Pt(0)
            r = p.add_run(str(value))
            set_font(r, size=font_size)
    if widths:
        for row in table.rows:
            for idx, width in enumerate(widths):
                row.cells[idx].width = Inches(width)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)
    return table


def add_code(doc, text):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    set_cell_shading(cell, "F6F8FA")
    set_cell_margins(cell, top=120, start=160, bottom=120, end=160)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    for idx, line in enumerate(text.splitlines()):
        if idx:
            p.add_run("\n")
        r = p.add_run(line)
        set_font(r, name="Consolas", size=8.5, color="24292F")
    doc.add_paragraph().paragraph_format.space_after = Pt(1)


def configure_document(doc):
    section = doc.sections[0]
    section.top_margin = Inches(0.85)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)
    section.header_distance = Inches(0.4)
    section.footer_distance = Inches(0.4)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    normal.paragraph_format.space_after = Pt(7)
    normal.paragraph_format.line_spacing = 1.25

    for name, size, color, before, after in (
        ("Heading 1", 16, BLUE, 16, 8),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 11.5, NAVY, 9, 4),
    ):
        style = styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for list_name in ("List Bullet", "List Bullet 2", "List Number"):
        style = styles[list_name]
        style.font.name = "Calibri"
        style.font.size = Pt(11)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.2

    header = section.header
    hp = header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = hp.add_run("Panduan Pengujian Model Prediksi Permintaan Bahan Kue")
    set_font(r, size=8.5, color=MUTED)

    footer = section.footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = fp.add_run("Tugas Akhir | Halaman ")
    set_font(r, size=9, color=MUTED)
    add_page_number(fp)


def add_cover(doc):
    for _ in range(5):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("PANDUAN PENGUJIAN MODEL")
    set_font(r, size=25, bold=True, color=NAVY)
    p.paragraph_format.space_after = Pt(6)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Linear Regression dan Random Forest")
    set_font(r, size=17, bold=True, color=BLUE)
    p.paragraph_format.space_after = Pt(8)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Studi Kasus Prediksi Permintaan Stok Bahan Kue")
    set_font(r, size=13, italic=True, color=MUTED)
    p.paragraph_format.space_after = Pt(34)

    add_table(
        doc,
        ["Informasi", "Keterangan"],
        [
            ["Jumlah data", "6.742 transaksi (6.741 data bersih)"],
            ["Pembagian data", "80% training dan 20% testing secara kronologis"],
            ["Target", "jumlah_permintaan_bahan"],
            ["Model yang dibandingkan", "Linear Regression dan Random Forest Regressor"],
            ["Tujuan dokumen", "Bahan belajar, pembahasan laporan, dan persiapan sidang"],
        ],
        widths=[1.8, 4.5],
        font_size=10,
    )

    for _ in range(5):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Disusun berdasarkan hasil pengujian model pada proyek tugas akhir")
    set_font(r, size=10, italic=True, color=MUTED)
    p.add_run().add_break(WD_BREAK.PAGE)


def build_document():
    doc = Document()
    configure_document(doc)
    add_cover(doc)

    add_heading(doc, "Ringkasan Eksekutif", 1)
    add_body(
        doc,
        "Pengujian model dilakukan untuk mengetahui kemampuan Linear Regression dan Random Forest "
        "dalam memprediksi jumlah permintaan bahan kue. Dataset berisi 6.742 transaksi, dengan 1 baris "
        "memiliki missing value pada produk_encoded sehingga menghasilkan 6.741 data bersih. Data tersebut "
        "dibagi secara kronologis menjadi 5.392 data training serta 1.349 data testing menggunakan rasio 80:20."
    )
    add_table(
        doc,
        ["Model", "R² Testing", "MAE", "RMSE", "Interpretasi Awal"],
        [
            ["Linear Regression", "0,7813", "0,63", "0,91", "Cukup baik pada test set"],
            ["Random Forest", "0,9964", "0,03", "0,12", "Sangat tinggi pada test set"],
        ],
        widths=[1.6, 1.0, 0.8, 0.8, 2.1],
    )
    add_callout(
        doc,
        "Temuan Kritis",
        "Nilai Random Forest yang sangat tinggi bukan sepenuhnya menunjukkan kemampuan meramalkan "
        "permintaan masa depan. Kolom total_harga_update memiliki hubungan langsung dengan target: "
        "total_harga_update = harga_satuan_update x jumlah_permintaan_bahan. Kondisi ini disebut data leakage.",
        fill=LIGHT_RED,
        title_color=RED,
    )
    add_body(
        doc,
        "Dokumen ini menjelaskan dua sudut pandang secara terpisah: pertama, bagaimana hasil testing "
        "yang ada dihitung; kedua, apakah hasil tersebut valid sebagai ukuran kemampuan prediksi nyata."
    )

    add_heading(doc, "Daftar Isi", 1)
    for item in (
        "1. Gambaran Umum Pengujian Model",
        "2. Persiapan dan Pembagian Data",
        "3. Cara Kerja Linear Regression",
        "4. Cara Kerja Random Forest",
        "5. Metrik Evaluasi Model",
        "6. Perhitungan Hasil Testing Proyek",
        "7. Perbandingan dan Interpretasi Model",
        "8. Analisis Data Leakage",
        "9. Pengujian Ulang Tanpa Data Leakage",
        "10. Rekomendasi Perbaikan Metodologi",
        "11. Contoh Penjelasan untuk Laporan",
        "12. Contoh Jawaban Saat Sidang",
        "Lampiran: Ringkasan Rumus dan Istilah",
    ):
        add_body(doc, item)
    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)

    add_heading(doc, "1. Gambaran Umum Pengujian Model", 1)
    add_body(
        doc,
        "Machine learning tidak cukup hanya dilatih. Model harus diuji menggunakan data yang tidak "
        "digunakan selama proses pelatihan. Tujuannya adalah mengukur apakah pola yang dipelajari model "
        "dapat diterapkan pada data lain, bukan hanya menghafal data training."
    )
    add_heading(doc, "1.1 Tujuan Testing", 2)
    for text in (
        "Menilai kedekatan hasil prediksi dengan nilai permintaan aktual.",
        "Membandingkan performa Linear Regression dan Random Forest.",
        "Mendeteksi overfitting, underfitting, dan kesalahan metodologi.",
        "Memilih model yang paling sesuai untuk digunakan pada aplikasi.",
    ):
        add_bullet(doc, text)
    add_heading(doc, "1.2 Alur Pengujian", 2)
    for text in (
        "Membaca dataset transaksi.",
        "Mengidentifikasi target dan fitur.",
        "Mengubah data kategori menjadi angka.",
        "Membagi data menjadi training set dan testing set.",
        "Melatih masing-masing model menggunakan training set.",
        "Menghasilkan prediksi pada testing set.",
        "Menghitung R², MAE, dan RMSE.",
        "Membandingkan serta mengevaluasi validitas hasil.",
    ):
        add_number(doc, text)

    add_heading(doc, "2. Persiapan dan Pembagian Data", 1)
    add_heading(doc, "2.1 Struktur Dataset", 2)
    add_table(
        doc,
        ["Komponen", "Nilai"],
        [
            ["Total data", "6.742 transaksi (6.741 data bersih)"],
            ["Training set", "5.392 transaksi atau 80% (periode 2021-01-01 s/d 2024-12-28)"],
            ["Testing set", "1.349 transaksi atau 20% (periode 2024-12-28 s/d 2025-12-31)"],
            ["Target", "jumlah_permintaan_bahan"],
            ["Jumlah fitur model tersimpan", "9 fitur"],
        ],
        widths=[2.3, 4.0],
    )
    add_heading(doc, "2.2 Fitur Model", 2)
    add_table(
        doc,
        ["No.", "Fitur", "Makna"],
        [
            ["1", "produk_encoded", "Kode angka produk (diubah menggunakan LabelEncoder)"],
            ["2", "tahun", "Tahun transaksi"],
            ["3", "bulan", "Bulan transaksi"],
            ["4", "hari", "Tanggal dalam bulan"],
            ["5", "hari_dalam_minggu", "Kode hari dalam minggu transaksi"],
            ["6", "harga_satuan_update", "Harga satuan bahan"],
            ["7", "hari_minggu", "Kode hari minggu (dayofweek)"],
            ["8", "nama_produk_encoded", "Nama produk yang telah diubah menggunakan LabelEncoder"],
            ["9", "kategori_produk_encoded", "Kategori produk yang telah diubah menggunakan LabelEncoder"],
        ],
        widths=[0.6, 2.3, 3.4],
    )
    add_heading(doc, "2.3 Pembagian Training dan Testing", 2)
    add_code(
        doc,
        "# Urutkan data secara kronologis\n"
        "df = df.sort_values(by='tanggal_transaksi')\n"
        "# Bagi data dengan rasio 80:20 secara manual\n"
        "split_index = int(len(X) * 0.8)\n"
        "X_train, X_test = X.iloc[:split_index], X.iloc[split_index:]\n"
        "y_train, y_test = y.iloc[:split_index], y.iloc[split_index:]"
    )
    add_body(
        doc,
        "Pembagian data dilakukan secara kronologis berdasarkan urutan tanggal transaksi. Parameter split_index "
        "diambil dari 80% panjang dataset bersih (6.741 data), yaitu int(6741 * 0.8) = 5.392 data training, "
        "dan sisa 1.349 data digunakan sebagai testing. Metode ini memastikan bahwa model diuji untuk memprediksi "
        "permintaan di masa depan, bukan sekadar memprediksi data acak di masa lalu."
    )
    add_callout(
        doc,
        "Pemberitahuan Metodologi",
        "Pengujian model menggunakan pembagian kronologis (chronological split) sangat penting untuk data deret "
        "waktu (time series) guna menghindari data leakage temporal, di mana model tidak sengaja mempelajari data masa depan.",
        fill=LIGHT_GOLD,
        title_color=GOLD,
    )

    add_heading(doc, "3. Cara Kerja Linear Regression", 1)
    add_body(
        doc,
        "Linear Regression mencari hubungan berbentuk garis lurus antara fitur masukan dan target. "
        "Setiap fitur memperoleh koefisien yang menunjukkan arah serta besarnya pengaruh terhadap hasil prediksi."
    )
    add_formula(doc, "ŷ = β₀ + β₁x₁ + β₂x₂ + ... + βₚxₚ", "Rumus umum prediksi Linear Regression")
    add_table(
        doc,
        ["Simbol", "Arti"],
        [
            ["ŷ", "Nilai jumlah permintaan hasil prediksi"],
            ["β₀", "Intercept atau nilai awal model"],
            ["βᵢ", "Koefisien untuk fitur ke-i"],
            ["xᵢ", "Nilai fitur ke-i"],
            ["p", "Jumlah fitur yang digunakan"],
        ],
        widths=[1.1, 5.2],
    )
    add_heading(doc, "3.1 Proses Pembelajaran", 2)
    add_body(
        doc,
        "Koefisien dicari dengan metode Ordinary Least Squares. Model memilih koefisien yang membuat "
        "jumlah kuadrat selisih antara nilai aktual dan prediksi menjadi sekecil mungkin."
    )
    add_formula(doc, "SSE = Σ(yᵢ - ŷᵢ)²", "SSE adalah jumlah kuadrat error prediksi")
    add_formula(doc, "β = (XᵀX)⁻¹Xᵀy", "Bentuk matriks solusi Ordinary Least Squares")
    add_heading(doc, "3.2 Persamaan yang Dihasilkan pada Proyek", 2)
    add_formula(
        doc,
        "ŷ ≈ 25,704 - 0,00306(produk) - 0,01028(tahun) + 0,00289(bulan) "
        "- 0,00276(hari) - 0,000208(harga satuan) + 0,0000421(total harga) + ...",
    )
    add_body(
        doc,
        "Linear Regression memperoleh R² testing sebesar 0,7813. Model ini terbantu oleh total harga, "
        "tetapi tidak dapat menangkap hubungan pembagian antara total harga dan harga satuan sebaik Random Forest."
    )
    add_heading(doc, "3.3 Kelebihan dan Keterbatasan", 2)
    add_table(
        doc,
        ["Kelebihan", "Keterbatasan"],
        [
            ["Mudah dijelaskan dan diinterpretasikan", "Mengasumsikan hubungan cenderung linear"],
            ["Proses training cepat", "Kurang baik untuk pola kompleks dan non-linear"],
            ["Koefisien menunjukkan arah hubungan", "Sensitif terhadap fitur yang tidak relevan dan outlier"],
        ],
        widths=[3.15, 3.15],
    )

    add_heading(doc, "4. Cara Kerja Random Forest", 1)
    add_body(
        doc,
        "Random Forest Regressor adalah metode ensemble yang membangun banyak Decision Tree. Setiap pohon "
        "mempelajari sampel dan kombinasi fitur yang berbeda. Prediksi akhir diperoleh dari rata-rata hasil seluruh pohon."
    )
    add_code(
        doc,
        "RandomForestRegressor(\n"
        "    n_estimators=100,\n"
        "    random_state=42,\n"
        "    n_jobs=-1\n"
        ")"
    )
    add_formula(doc, "ŷRF = (ŷ₁ + ŷ₂ + ... + ŷT) / T", "Prediksi Random Forest adalah rata-rata prediksi T pohon")
    add_body(doc, "Karena n_estimators=100, nilai T pada proyek ini adalah 100 pohon.")
    add_heading(doc, "4.1 Proses pada Setiap Pohon", 2)
    for text in (
        "Mengambil sampel data training secara acak dengan pengembalian atau bootstrap.",
        "Memilih sebagian fitur secara acak pada setiap kandidat percabangan.",
        "Mencari batas pembagian data yang paling mengurangi error.",
        "Menghasilkan prediksi dari nilai rata-rata target pada leaf node.",
        "Menggabungkan prediksi seluruh pohon dengan operasi rata-rata.",
    ):
        add_number(doc, text)
    add_heading(doc, "4.2 Pemilihan Percabangan", 2)
    add_body(
        doc,
        "Untuk regresi, Decision Tree memilih percabangan yang menurunkan Mean Squared Error. "
        "Percabangan yang membuat nilai target dalam setiap kelompok semakin seragam akan dipilih."
    )
    add_formula(doc, "MSE = (1/n) Σ(yᵢ - ŷᵢ)²")
    add_heading(doc, "4.3 Mengapa Random Forest Lebih Fleksibel", 2)
    add_body(
        doc,
        "Berbeda dari Linear Regression, Random Forest tidak mengharuskan hubungan berbentuk garis lurus. "
        "Model dapat membentuk banyak aturan kondisi sehingga mampu mempelajari interaksi kompleks antara fitur."
    )

    add_heading(doc, "5. Metrik Evaluasi Model", 1)
    add_heading(doc, "5.1 R² atau Koefisien Determinasi", 2)
    add_formula(doc, "R² = 1 - [Σ(yᵢ - ŷᵢ)² / Σ(yᵢ - ȳ)²]")
    add_body(
        doc,
        "R² mengukur seberapa besar variasi target yang dapat dijelaskan model. Nilai mendekati 1 "
        "menunjukkan prediksi semakin dekat dengan data aktual. Nilai 0 berarti performa setara dengan "
        "selalu menebak rata-rata. Nilai negatif berarti model lebih buruk daripada tebakan rata-rata."
    )
    add_callout(
        doc,
        "Istilah yang Tepat",
        "R² sering disebut sebagai akurasi pada laporan regresi. Istilah akademik yang lebih tepat adalah "
        "koefisien determinasi atau kemampuan model menjelaskan variasi target.",
    )
    add_heading(doc, "5.2 Mean Absolute Error", 2)
    add_formula(doc, "MAE = (1/n) Σ|yᵢ - ŷᵢ|")
    add_body(
        doc,
        "MAE menunjukkan rata-rata besar kesalahan prediksi tanpa memperhatikan arah kesalahan. "
        "Jika MAE sebesar 0,63 unit, prediksi rata-rata meleset sekitar 0,63 unit dari nilai aktual."
    )
    add_heading(doc, "5.3 Root Mean Squared Error", 2)
    add_formula(doc, "RMSE = √[(1/n) Σ(yᵢ - ŷᵢ)²]")
    add_body(
        doc,
        "RMSE memberikan penalti lebih besar pada kesalahan yang besar karena error dikuadratkan sebelum "
        "dirata-ratakan. Semakin kecil nilai MAE dan RMSE, semakin baik kedekatan prediksi."
    )
    add_heading(doc, "5.4 Ringkasan Interpretasi", 2)
    add_table(
        doc,
        ["Metrik", "Arah Nilai yang Baik", "Makna"],
        [
            ["R²", "Semakin mendekati 1", "Besarnya variasi target yang dijelaskan model"],
            ["MAE", "Semakin mendekati 0", "Rata-rata kesalahan absolut"],
            ["RMSE", "Semakin mendekati 0", "Kesalahan dengan penalti lebih besar untuk error besar"],
        ],
        widths=[1.0, 1.8, 3.5],
    )

    add_heading(doc, "6. Perhitungan Hasil Testing Proyek", 1)
    add_body(
        doc,
        "Pada testing set terdapat 1.349 data. Rata-rata target testing adalah sekitar 5,4188 unit, "
        "sedangkan total variasi aktual atau SST adalah 10.964,3617."
    )
    add_heading(doc, "6.1 Linear Regression", 2)
    add_table(
        doc,
        ["Komponen", "Nilai"],
        [
            ["Jumlah absolute error atau SAE", "3.357,4719"],
            ["Jumlah squared error atau SSE", "11.007,5413"],
            ["Jumlah data testing", "1.349"],
            ["Total variasi aktual atau SST", "10.964,3617"],
        ],
        widths=[3.4, 2.9],
    )
    add_formula(doc, "MAE = 3.357,4719 / 1.349 = 2,4889 ≈ 2,49")
    add_formula(doc, "RMSE = √(11.007,5413 / 1.349) = 2,8565 ≈ 2,86")
    add_formula(doc, "R² = 1 - (11.007,5413 / 10.964,3617) = -0,0039")
    add_body(
        doc,
        "Interpretasinya, Linear Regression tidak mampu menjelaskan variasi target pada testing set secara optimal "
        "(R² bernilai negatif sebesar -0,0039 atau -0,39%) dan rata-rata prediksinya meleset sekitar 2,49 unit."
    )
    add_heading(doc, "6.2 Random Forest", 2)
    add_table(
        doc,
        ["Komponen", "Nilai"],
        [
            ["Jumlah absolute error atau SAE", "3.494,5478"],
            ["Jumlah squared error atau SSE", "12.431,3712"],
            ["Jumlah data testing", "1.349"],
            ["Total variasi aktual atau SST", "10.964,3617"],
        ],
        widths=[3.4, 2.9],
    )
    add_formula(doc, "MAE = 3.494,5478 / 1.349 = 2,5905 ≈ 2,59")
    add_formula(doc, "RMSE = √(12.431,3712 / 1.349) = 3,0357 ≈ 3,04")
    add_formula(doc, "R² = 1 - (12.431,3712 / 10.964,3617) = -0,1338")
    add_body(
        doc,
        "Interpretasinya, Random Forest tidak mampu menjelaskan variasi target pada testing set secara optimal "
        "(R² bernilai negatif sebesar -0,1338 atau -13,38%) dan rata-rata prediksinya meleset sekitar 2,59 unit "
        "setelah fitur total_harga_update yang menimbulkan data leakage dihapus."
    )
    add_heading(doc, "6.3 Contoh Hasil Prediksi Individual", 2)
    add_table(
        doc,
        ["Aktual", "Prediksi Linear Regression", "Prediksi Random Forest", "Error Absolut RF"],
        [
            ["7", "5,4795", "3,71", "3,29"],
            ["5", "5,5121", "4,99", "0,01"],
            ["10", "5,5587", "5,04", "4,96"],
            ["1", "5,4686", "5,67", "4,67"],
            ["2", "5,4883", "5,64", "3,64"],
        ],
        widths=[0.9, 2.0, 2.0, 1.4],
        font_size=9,
    )

    add_heading(doc, "7. Perbandingan dan Interpretasi Model", 1)
    add_table(
        doc,
        ["Aspek", "Linear Regression", "Random Forest"],
        [
            ["R² testing", "-0,0039", "-0,1338"],
            ["MAE testing", "2,49 unit", "2,59 unit"],
            ["RMSE testing", "2,86 unit", "3,04 unit"],
            ["Kemampuan pola non-linear", "Terbatas", "Baik (pada data training)"],
            ["Kemudahan interpretasi", "Sangat mudah", "Lebih kompleks"],
            ["Ketergantungan data leakage", "Bebas leakage (total harga dihapus)", "Bebas leakage (total harga dihapus)"],
        ],
        widths=[2.2, 2.05, 2.05],
    )
    add_heading(doc, "7.1 Perbandingan MAE", 2)
    add_formula(doc, "Peningkatan MAE = [(2,5905 - 2,4889) / 2,4889] x 100% = 4,08%")
    add_body(
        doc,
        "Pernyataan yang tepat adalah model Random Forest menghasilkan kesalahan rata-rata yang sedikit "
        "lebih tinggi (+4,08%) dibandingkan Linear Regression pada test set aktual tanpa leakage."
    )
    add_heading(doc, "7.2 Indikasi Overfitting", 2)
    add_table(
        doc,
        ["Model", "R² Training", "R² Testing", "Selisih"],
        [
            ["Linear Regression", "0,0013", "-0,0039", "0,0052"],
            ["Random Forest", "0,6913", "-0,1338", "0,8251"],
        ],
        widths=[2.0, 1.4, 1.4, 1.4],
    )
    add_body(
        doc,
        "Selisih training dan testing pada Linear Regression terlihat sangat kecil, namun kedua nilainya mendekati nol, "
        "menunjukkan underfitting. Sedangkan pada Random Forest, terdapat selisih R² yang sangat besar (0,8251) "
        "antara training (0,6913) dan testing (-0,1338). Hal ini menunjukkan indikasi overfitting yang sangat kuat "
        "pada Random Forest setelah total harga dihapus dari fitur. Model berkinerja cukup baik pada data training "
        "namun gagal memprediksi data testing baru yang belum pernah dilihat sebelumnya."
    )

    add_heading(doc, "8. Analisis Data Leakage", 1)
    add_body(
        doc,
        "Data leakage adalah kondisi ketika fitur model mengandung informasi yang seharusnya belum "
        "tersedia pada saat prediksi dilakukan atau secara langsung dibentuk dari target. Leakage membuat "
        "hasil testing terlihat sangat baik tetapi performa tersebut sulit dipertahankan pada penggunaan nyata."
    )
    add_formula(doc, "total_harga_update = harga_satuan_update x jumlah_permintaan_bahan")
    add_formula(doc, "jumlah_permintaan_bahan = total_harga_update / harga_satuan_update")
    add_callout(
        doc,
        "Hasil Pemeriksaan Dataset",
        "Hubungan total_harga_update = harga_satuan_update x jumlah_permintaan_bahan berlaku tepat pada "
        "seluruh 6.742 baris atau 100% data.",
        fill=LIGHT_RED,
        title_color=RED,
    )
    add_heading(doc, "8.1 Bukti dari Feature Importance", 2)
    add_table(
        doc,
        ["Fitur", "Feature Importance Random Forest"],
        [
            ["total_harga_update", "62,97%"],
            ["harga_satuan_update", "36,18%"],
            ["Gabungan kedua fitur harga", "99,15%"],
            ["Seluruh fitur lain", "Kurang dari 1% secara individual"],
        ],
        widths=[3.7, 2.6],
    )
    add_body(
        doc,
        "Feature importance menunjukkan bahwa Random Forest hampir sepenuhnya menggunakan total harga "
        "dan harga satuan. Model terutama mempelajari cara memperoleh kembali jumlah permintaan, bukan "
        "mempelajari pola permintaan berdasarkan waktu, produk, atau kategori."
    )
    add_heading(doc, "8.2 Mengapa Leakage Berbahaya", 2)
    for text in (
        "Total harga transaksi biasanya baru diketahui setelah jumlah permintaan diketahui.",
        "Pada saat memprediksi permintaan masa depan, total harga belum tersedia.",
        "Nilai R² yang tinggi dapat memberikan kesimpulan keliru bahwa model siap digunakan.",
        "Model berpotensi gagal ketika menerima data nyata tanpa informasi target terselubung.",
    ):
        add_bullet(doc, text)

    add_heading(doc, "9. Pengujian Ulang Tanpa Data Leakage", 1)
    add_body(
        doc,
        "Untuk mengukur kemampuan prediksi yang lebih realistis, total_harga_update dihapus dari fitur. "
        "Pengujian dilakukan menggunakan random split dan chronological split."
    )
    add_table(
        doc,
        ["Skenario", "Model", "R²", "MAE", "RMSE"],
        [
            ["Chronological split tanpa total harga (backend Railway)", "Linear Regression", "-0,0039", "2,49", "2,86"],
            ["Chronological split tanpa total harga (backend Railway)", "Random Forest", "-0,1338", "2,59", "3,04"],
        ],
        widths=[2.3, 1.6, 0.8, 0.8, 0.8],
        font_size=8.8,
    )
    add_callout(
        doc,
        "Makna R² Negatif",
        "R² negatif berarti model lebih buruk dibandingkan metode sederhana yang selalu memprediksi "
        "nilai rata-rata permintaan. Hasil ini menunjukkan fitur yang tersedia belum mengandung pola "
        "permintaan yang cukup kuat setelah total harga dihapus.",
        fill=LIGHT_GOLD,
        title_color=GOLD,
    )
    add_heading(doc, "9.1 Kesimpulan Validitas", 2)
    add_body(
        doc,
        "Angka R² sebesar 0,9964 tetap merupakan hasil perhitungan yang benar untuk konfigurasi testing "
        "saat ini. Akan tetapi, angka tersebut tidak valid jika digunakan untuk menyatakan kemampuan "
        "memprediksi permintaan masa depan secara murni karena terdapat data leakage."
    )

    add_heading(doc, "10. Rekomendasi Perbaikan Metodologi", 1)
    add_heading(doc, "10.1 Perbaikan Fitur", 2)
    for text in (
        "Hapus total_harga_update dari fitur prediksi.",
        "Hapus salah satu dari hari_dalam_minggu atau hari_minggu karena keduanya duplikat.",
        "Gunakan lag permintaan, misalnya jumlah permintaan hari sebelumnya atau minggu sebelumnya.",
        "Gunakan moving average permintaan dalam 7, 14, atau 30 hari terakhir.",
        "Tambahkan informasi hari libur, musim, promosi, dan periode ramai.",
        "Tambahkan stok sebelumnya serta histori pemakaian bahan.",
    ):
        add_bullet(doc, text)
    add_heading(doc, "10.2 Perbaikan Proses Testing", 2)
    for text in (
        "Gunakan chronological split: data lama untuk training dan data terbaru untuk testing.",
        "Gunakan TimeSeriesSplit atau walk-forward validation untuk evaluasi yang lebih kuat.",
        "Bandingkan model dengan baseline sederhana seperti rata-rata atau nilai periode sebelumnya.",
        "Lakukan tuning parameter Random Forest hanya menggunakan data training.",
        "Laporkan R², MAE, dan RMSE secara bersamaan.",
    ):
        add_bullet(doc, text)
    add_heading(doc, "10.3 Konsistensi Aplikasi Flutter", 2)
    add_body(
        doc,
        "Python menggunakan kode hari 0 sampai 6 melalui dt.dayofweek (Senin=0), sedangkan DateTime.weekday pada Flutter menggunakan "
        "1 sampai 7 (Senin=1). Nilai harus dipetakan secara konsisten sebelum dikirim ke backend Railway. Selain itu, pemetaan encoded produk "
        "dan kategori pada Flutter harus sama persis dengan LabelEncoder saat training di backend Flask."
    )

    add_heading(doc, "11. Contoh Penjelasan untuk Laporan", 1)
    add_heading(doc, "11.1 Paragraf Metode Pengujian", 2)
    add_callout(
        doc,
        "Contoh Penulisan",
        "Pengujian model dilakukan menggunakan metode chronological split dengan membagi dataset bersih "
        "menjadi 80% data training dan 20% data testing. Dari total 6.742 data transaksi (6.741 data bersih setelah pembersihan 1 missing value), sebanyak 5.392 data digunakan "
        "sebagai training set (periode awal 2021-01-01 s/d 2024-12-28) dan 1.349 data digunakan sebagai testing set "
        "(periode akhir 2024-12-28 s/d 2025-12-31). Pembagian data diurutkan secara kronologis berdasarkan tanggal transaksi "
        "untuk mengevaluasi kemampuan peramalan model secara realistis. Evaluasi dilakukan menggunakan "
        "koefisien determinasi R², Mean Absolute Error, dan Root Mean Squared Error.",
    )
    add_heading(doc, "11.2 Paragraf Hasil Perbandingan", 2)
    add_callout(
        doc,
        "Contoh Penulisan",
        "Berdasarkan hasil pengujian model tanpa data leakage, Linear Regression menghasilkan R² sebesar -0,0039, "
        "MAE sebesar 2,49, dan RMSE sebesar 2,86. Random Forest menghasilkan R² sebesar -0,1338, MAE sebesar 2,59, "
        "dan RMSE sebesar 3,04. Secara numerik, model Linear Regression memberikan kesalahan rata-rata (MAE) "
        "yang sedikit lebih rendah dibandingkan Random Forest pada data testing.",
    )
    add_heading(doc, "11.3 Paragraf Keterbatasan", 2)
    add_callout(
        doc,
        "Contoh Penulisan",
        "Evaluasi model ini dilakukan setelah mengeluarkan fitur total_harga_update untuk menghindari data leakage. "
        "Pada pengujian awal sebelum fitur tersebut dikeluarkan, model Random Forest sempat memperoleh R² sebesar 0,9964, "
        "MAE sebesar 0,03, dan RMSE sebesar 0,12. Namun, nilai tersebut tidak valid untuk menggambarkan kemampuan "
        "prediksi masa depan karena fitur total harga dihitung langsung dari target (jumlah permintaan). Setelah perbaikan "
        "metodologi dengan menghapus fitur total harga, diperoleh hasil pengujian aktual yang lebih realistis dan aman dari data leakage.",
        fill=LIGHT_GOLD,
        title_color=GOLD,
    )
    add_heading(doc, "11.4 Paragraf Rekomendasi Pengembangan", 2)
    add_callout(
        doc,
        "Contoh Penulisan",
        "Pengembangan berikutnya perlu menghapus fitur total harga dan menggunakan fitur yang telah tersedia "
        "sebelum transaksi terjadi, seperti histori permintaan, lag permintaan, moving average, tren bulanan, "
        "hari libur, promosi, dan stok sebelumnya. Pengujian juga perlu dilakukan secara kronologis agar "
        "lebih sesuai dengan tujuan prediksi masa depan.",
        fill="ECF7EF",
        title_color=GREEN,
    )

    add_heading(doc, "12. Contoh Jawaban Saat Sidang", 1)
    add_heading(doc, "12.1 Bagaimana Proses Testing Dilakukan?", 2)
    add_body(
        doc,
        "Jawaban: Proses testing dilakukan dengan membagi 6.742 data transaksi (dengan 1 baris dieliminasi karena "
        "mengandung missing value sehingga tersisa 6.741 data bersih) menjadi 80% data training (5.392 data, periode "
        "awal) dan 20% data testing (1.349 data, periode akhir) secara kronologis berdasarkan urutan tanggal. Model "
        "mempelajari pola dari data training untuk kemudian memprediksi data testing. Hasil prediksi "
        "dievaluasi menggunakan R², MAE, dan RMSE."
    )
    add_heading(doc, "12.2 Mengapa Menggunakan Tiga Metrik?", 2)
    add_body(
        doc,
        "Jawaban: R² digunakan untuk mengukur seberapa besar variasi target dapat dijelaskan model. MAE "
        "menunjukkan rata-rata kesalahan dalam satuan permintaan sehingga mudah dipahami. RMSE memberikan "
        "penalti lebih besar untuk kesalahan yang besar. Ketiganya memberikan penilaian yang lebih lengkap."
    )
    add_heading(doc, "12.3 Mengapa Hasil Evaluasi R² Bernilai Negatif?", 2)
    add_body(
        doc,
        "Jawaban: Nilai R² negatif pada data testing (Linear Regression: -0,0039, Random Forest: -0,1338) menunjukkan "
        "bahwa model kesulitan untuk memprediksi data masa depan berdasarkan urutan waktu (chronological split) setelah "
        "fitur data leakage (total_harga_update) dihapus. Hal ini mengonfirmasi bahwa fitur waktu kalender saja belum cukup "
        "kuat untuk menangkap fluktuasi pola permintaan rill, sehingga performanya di bawah tebakan rata-rata."
    )
    add_heading(doc, "12.4 Bagaimana Hasil Model yang Di-deploy di Railway?", 2)
    add_body(
        doc,
        "Jawaban: Model backend Flask yang di-deploy di Railway adalah model Random Forest yang dilatih tanpa menggunakan fitur total_harga_update "
        "untuk memastikan kebebasan dari data leakage. Model ini diakses oleh aplikasi Flutter sebagai frontend melalui API. Model ini memiliki metrik evaluasi R² sebesar -0,1338, MAE sebesar 2,59, "
        "dan RMSE sebesar 3,04 pada data testing. Walaupun secara angka lebih rendah dibandingkan model awal dengan leakage, "
        "performa ini jauh lebih valid dan dapat diandalkan secara akademis."
    )
    add_heading(doc, "12.5 Apakah Model Mengalami Overfitting Setelah Leakage Dihapus?", 2)
    add_body(
        doc,
        "Jawaban: Ya. Setelah fitur total_harga_update dihapus, model Random Forest menunjukkan indikasi overfitting yang sangat kuat "
        "di mana R² training bernilai cukup baik (0,6913) tetapi R² testing anjlok menjadi negatif (-0,1338). Ini membuktikan model "
        "menghafal pola data training tetapi gagal dalam melakukan generalisasi pada data pengujian baru."
    )
    add_heading(doc, "12.6 Apa Perbaikan yang Harus Dilakukan?", 2)
    add_body(
        doc,
        "Jawaban: Total harga harus dihapus dari fitur, kemudian model dilatih menggunakan fitur yang benar-benar "
        "tersedia sebelum transaksi terjadi. Testing sebaiknya menggunakan pembagian berdasarkan waktu dan "
        "dibandingkan dengan baseline. Fitur histori seperti lag dan moving average juga perlu ditambahkan."
    )

    add_heading(doc, "Lampiran: Ringkasan Rumus dan Istilah", 1)
    add_table(
        doc,
        ["Konsep", "Rumus atau Definisi Ringkas"],
        [
            ["Linear Regression", "ŷ = β₀ + β₁x₁ + ... + βₚxₚ"],
            ["Random Forest", "Rata-rata prediksi dari banyak Decision Tree"],
            ["SSE", "Σ(yᵢ - ŷᵢ)²"],
            ["MAE", "(1/n) Σ|yᵢ - ŷᵢ|"],
            ["RMSE", "√[(1/n) Σ(yᵢ - ŷᵢ)²]"],
            ["R²", "1 - SSE/SST"],
            ["Overfitting", "Model sangat baik pada training tetapi buruk pada data baru"],
            ["Underfitting", "Model terlalu sederhana dan gagal mempelajari pola"],
            ["Data leakage", "Fitur mengandung informasi target atau informasi masa depan"],
            ["Chronological split", "Training memakai data lama, testing memakai data terbaru"],
            ["Baseline", "Metode sederhana sebagai pembanding minimum"],
        ],
        widths=[2.0, 4.3],
        font_size=9.3,
    )
    add_callout(
        doc,
        "Kesimpulan Akhir",
        "Random Forest menghasilkan metrik terbaik pada testing awal, tetapi performa tersebut sangat "
        "dipengaruhi data leakage dari total_harga_update. Untuk laporan yang bertanggung jawab secara "
        "akademik, tampilkan hasil awal, jelaskan keterbatasannya, dan cantumkan rekomendasi pengujian ulang "
        "tanpa fitur yang dibentuk dari target.",
        fill="ECF7EF",
        title_color=GREEN,
    )

    path1 = os.path.join(script_dir, 'Panduan_Pengujian_Linear_Regression_dan_Random_Forest.docx')
    path2 = os.path.join(os.path.dirname(script_dir), 'Panduan_Pengujian_Linear_Regression_dan_Random_Forest.docx')
    
    try:
        doc.save(path1)
        print(f"Saved: {path1}")
    except Exception as e:
        print(f"Error saving to {path1}: {e}")
        
    try:
        doc.save(path2)
        print(f"Saved: {path2}")
    except Exception as e:
        backup_path = os.path.join(os.path.dirname(script_dir), 'Panduan_Pengujian_Linear_Regression_dan_Random_Forest_Project.docx')
        doc.save(backup_path)
        print(f"Warning: File {path2} is locked. Saved as backup: {backup_path}")


if __name__ == "__main__":
    build_document()
